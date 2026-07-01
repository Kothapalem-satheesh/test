from docx import Document
import io
from app.services.chunker import chunk_text, detect_section_title


def parse_docx(file_bytes: bytes) -> tuple[list[dict], int]:
    """Extract text from DOCX with section awareness."""
    doc = Document(io.BytesIO(file_bytes))
    chunks = []
    current_section = None
    current_text = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect headings by style
        if para.style and para.style.name and "Heading" in para.style.name:
            if current_text:
                section_chunks = chunk_text(
                    current_text,
                    section_title=current_section,
                )
                chunks.extend(section_chunks)
                current_text = ""
            current_section = text
        else:
            current_text += f"\n\n{text}" if current_text else text

    if current_text:
        section_chunks = chunk_text(current_text, section_title=current_section)
        chunks.extend(section_chunks)

    # Extract tables
    for table in doc.tables:
        table_text = "\n".join(
            " | ".join(cell.text.strip() for cell in row.cells)
            for row in table.rows
        )
        if table_text.strip():
            table_chunks = chunk_text(
                f"[Table]\n{table_text}",
                section_title=current_section or "Table Data",
            )
            chunks.extend(table_chunks)

    for i, chunk in enumerate(chunks):
        chunk["chunk_index"] = i

    return chunks, len(doc.paragraphs)
